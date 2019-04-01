#! python3

''' This code takes a txt file with names as input and writes
    separate custom invites for each of the names
    TODO: Add Style. Couldnt get styles to work.
'''

from docx import Document

#Take the input the name file and open this as well as the new DOCX

file=open(input('Enter the name of the user list file ... '))
doc=Document('Stuff.docx')

#TODO Loop through each line and create the custom Word page

for line in file.readlines():
    doc.add_paragraph('It would be a pleasure to have the company of')
    doc.add_paragraph(line)
    doc.add_paragraph('at 1101101 Memory Lane on the evening of')
    doc.add_paragraph('April 1st')
    doc.add_paragraph("7 O'Clock")
    doc.add_page_break()
    
#Save to the previously created invites.docx to use the custom style    

doc.save('Stuff.docx')    
